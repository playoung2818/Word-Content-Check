from __future__ import annotations

from collections import Counter, defaultdict
from datetime import datetime, timedelta
from pathlib import Path
import os
import re
from typing import Any

import gspread
import pandas as pd
from docx import Document
from gspread_dataframe import get_as_dataframe
from sqlalchemy import create_engine, text


DB_URL = os.getenv("WO_DB_URL", "postgresql://postgres:Czheyuan0227%40@192.168.60.215:5432/postgres")
GOOGLE_CRED_PATH = Path(os.getenv("WO_GOOGLE_CRED_PATH", r"D:\Python\pdfwo-466115-734096e1cef8.json"))
RECEIVING_LOG_PATH = Path(os.getenv("WO_RECEIVING_LOG_PATH", r"D:\OneDrive - neousys-tech\Share NTA Warehouse\01 Incoming\Receiving Log_ZC_2.0.xlsm"))
WO_FOLDER = Path(os.getenv("WO_FOLDER", r"D:\OneDrive - neousys-tech\Share NTA Warehouse\02 Work Order- Word file\Work Order 2026\Work Order 2026-06"))

SHEET_NAME = os.getenv("WO_SHEET_NAME", "PDF_WO")
SALES_ORDER_TAB = os.getenv("WO_SALES_ORDER_TAB", "Open Sales Order")
RED_ALERT = "\033[1;97;41m CRITICAL \033[0m"
YELLOW_ALERT = "\033[1;30;43m CHECK \033[0m"


def make_engine(db_url: str | None = None):
    db_url = db_url or DB_URL
    if not db_url:
        raise ValueError("Set DB_URL or the WO_DB_URL environment variable first.")
    return create_engine(db_url)


def clean_text(value: Any) -> str:
    if pd.isna(value):
        return ""
    return str(value).strip()


def normalize_part(value: Any) -> str:
    return clean_text(value).upper().replace("-", "").replace(" ", "")


def read_sales_order(
    google_cred_path: Path | str = GOOGLE_CRED_PATH,
    sheet_name: str = SHEET_NAME,
    sales_order_tab: str = SALES_ORDER_TAB,
) -> pd.DataFrame:
    gc = gspread.service_account(filename=str(google_cred_path))
    ws = gc.open(sheet_name).worksheet(sales_order_tab)
    df = get_as_dataframe(ws, evaluate_formulas=True, dtype=str).dropna(how="all")
    if "QB Num" not in df.columns:
        raise ValueError("Expected 'QB Num' column not found in Open Sales Order sheet.")
    df["QB Num"] = df["QB Num"].map(clean_text)
    return df


def clean_receiving_log(file_path: Path | str = RECEIVING_LOG_PATH) -> pd.DataFrame:
    columns = {
        "Date": "entry_date",
        "Inv# ": "invoice_number",
        "Box #": "box_number",
        "POD#": "pod_number",
        "Part#": "part_number",
        "SN#": "serial_number",
        "QTY": "quantity",
    }
    df = pd.read_excel(file_path, sheet_name="Receiving").rename(columns=columns)
    df = df.dropna(subset=list(columns.values()), how="all")
    df["quantity"] = pd.to_numeric(df["quantity"], errors="coerce").fillna(1).astype(float)
    df["entry_date"] = pd.to_datetime(df["entry_date"], errors="coerce")

    for col in ["invoice_number", "box_number", "pod_number", "part_number", "serial_number"]:
        df[col] = df[col].map(clean_text)

    return df[df["serial_number"] != ""].copy()


def load_receiving_log(file_path: Path | str = RECEIVING_LOG_PATH, dry_run: bool = True) -> pd.DataFrame:
    df = clean_receiving_log(file_path)
    key_cols = ["entry_date", "invoice_number", "box_number", "pod_number", "part_number", "serial_number", "quantity"]

    engine = make_engine()
    existing = pd.read_sql(f"SELECT {', '.join(key_cols)} FROM receiving_log", engine)
    existing["entry_date"] = pd.to_datetime(existing["entry_date"], errors="coerce")
    existing["quantity"] = pd.to_numeric(existing["quantity"], errors="coerce").astype(float)
    for col in ["invoice_number", "box_number", "part_number", "serial_number", "pod_number"]:
        existing[col] = existing[col].map(clean_text)

    new_rows = df.merge(existing, on=key_cols, how="left", indicator=True)
    new_rows = new_rows[new_rows["_merge"] == "left_only"].drop(columns="_merge")

    print(f"{len(new_rows)} new rows found out of {len(df)} cleaned rows.")
    if dry_run:
        return new_rows

    new_rows.to_sql("receiving_log", engine, if_exists="append", index=False, method="multi")
    print("Inserted new receiving_log rows.")
    return new_rows


def extract_wo_number(file_path: Path | str) -> str | None:
    match = re.search(r"WO\d{0,2}-(\d{8})", Path(file_path).name, flags=re.I)
    return f"SO-{match.group(1)}" if match else None


def is_real_wo_file(file_path: Path | str) -> bool:
    return extract_wo_number(file_path) is not None


def extract_product_details(file_path: Path | str) -> list[dict[str, Any]]:
    document = Document(file_path)
    if not document.tables:
        return []

    rows = []
    for row_index, row in enumerate(document.tables[0].rows[1:-1], start=1):
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 4:
            rows.append({
                "row_index": row_index,
                "product_number": cells[0] if len(cells) > 0 else "",
                "qty": cells[1] if len(cells) > 1 else "",
                "serials": [],
                "notes": "",
                "row_warning": "Expected at least 4 columns in the WO table.",
            })
            continue

        rows.append({
            "row_index": row_index,
            "product_number": cells[0],
            "qty": cells[1],
            "serials": [
                s.strip()
                for s in cells[2].splitlines()
                if s.strip() and s.strip().upper() not in {"NA", "N/A", "NONE"}
            ],
            "notes": cells[3],
            "row_warning": "",
        })
    return rows


def compare_word_file_to_sales_order_result(file_path: Path | str, sales_order: pd.DataFrame | None = None) -> dict[str, Any]:
    word_items = extract_product_details(file_path)
    word_count = len(word_items)
    wo_number = extract_wo_number(file_path)
    result = {
        "file": Path(file_path).name,
        "wo_number": wo_number,
        "severity": "INFO",
        "error_type": "ROW_COUNT_OK",
        "word_count": word_count,
        "sheet_count": None,
        "message": "Open Sales Order row count and item names OK",
        "item_mismatches": [],
    }

    if sales_order is None:
        result.update({
            "severity": "WARNING",
            "error_type": "SALES_ORDER_NOT_LOADED",
            "message": f"Word rows {word_count}; Google Sheet not loaded",
        })
        return result
    if not wo_number:
        result.update({"severity": "WARNING", "error_type": "MISSING_WO_NUMBER", "message": "No WO number found in filename"})
        return result

    lookup_col = "QB Num" if "QB Num" in sales_order.columns else "WO_Number" if "WO_Number" in sales_order.columns else None
    if lookup_col is None:
        result.update({
            "severity": "WARNING",
            "error_type": "SALES_ORDER_COLUMN_MISSING",
            "message": "Missing QB Num or WO_Number column in Open Sales Order",
        })
        return result

    sheet_rows = sales_order[sales_order[lookup_col].map(clean_text) == wo_number].copy()
    sheet_count = len(sheet_rows)
    result["sheet_count"] = sheet_count
    if word_count != sheet_count:
        result.update({
            "severity": "CRITICAL",
            "error_type": "ROW_COUNT_MISMATCH",
            "message": f"Row count mismatch: Word {word_count}, Open Sales Order {sheet_count}",
        })
    if "Item" not in sheet_rows.columns:
        if result["severity"] != "CRITICAL":
            result.update({"severity": "WARNING", "error_type": "SALES_ORDER_ITEM_COLUMN_MISSING"})
        result["message"] += "; missing Item column in Open Sales Order"
        return result

    word_parts = [clean_text(item.get("product_number", "")) for item in word_items]
    sheet_parts = [clean_text(item) for item in sheet_rows["Item"].tolist()]
    word_counts = Counter(normalize_part(part) for part in word_parts)
    sheet_counts = Counter(normalize_part(part) for part in sheet_parts)
    word_names: dict[str, list[str]] = defaultdict(list)
    sheet_names: dict[str, list[str]] = defaultdict(list)
    for part in word_parts:
        word_names[normalize_part(part)].append(part)
    for part in sheet_parts:
        sheet_names[normalize_part(part)].append(part)

    item_mismatches = []
    for part_key, count in (word_counts - sheet_counts).items():
        item_mismatches.append({
            "mismatch_type": "MISSING_IN_SHEET",
            "word_part": word_names[part_key][0] if word_names[part_key] else "",
            "sheet_part": "",
            "count": count,
        })
    for part_key, count in (sheet_counts - word_counts).items():
        item_mismatches.append({
            "mismatch_type": "EXTRA_IN_SHEET",
            "word_part": "",
            "sheet_part": sheet_names[part_key][0] if sheet_names[part_key] else "",
            "count": count,
        })

    if item_mismatches:
        result["item_mismatches"] = item_mismatches
        if result["error_type"] == "ROW_COUNT_MISMATCH":
            result["message"] += f"; {len(item_mismatches)} item name mismatch(es)"
        else:
            result.update({
                "severity": "CRITICAL",
                "error_type": "ITEM_NAME_MISMATCH",
                "message": f"Item name mismatch: Word and Open Sales Order have different item set ({len(item_mismatches)} difference(s))",
            })
    return result


def sales_order_row_count_result(file_path: Path | str, sales_order: pd.DataFrame | None = None) -> dict[str, Any]:
    return compare_word_file_to_sales_order_result(file_path, sales_order=sales_order)


def compare_word_file_to_sales_order(file_path: Path | str, sales_order: pd.DataFrame | None = None) -> str:
    result = compare_word_file_to_sales_order_result(file_path, sales_order=sales_order)
    if result["severity"] == "CRITICAL":
        status = "CRITICAL"
    elif result["severity"] == "WARNING":
        status = f"WARNING {result['error_type']}"
    else:
        status = "OK"
    item_mismatch_count = len(result.get("item_mismatches", []))
    return (
        f"Status: {status} | Open Sales Order row count: Word {result['word_count']}, "
        f"Google Sheet {result['sheet_count']} | Item mismatches: {item_mismatch_count}"
    )


def db_part_for_serial(serial_number: str) -> str | None:
    engine = make_engine()
    with engine.begin() as conn:
        return conn.execute(
            text("SELECT part_number FROM receiving_log WHERE serial_number = :sn ORDER BY entry_date DESC, id DESC LIMIT 1"),
            {"sn": serial_number},
        ).scalar_one_or_none()


def make_issue(
    file: Path | str,
    wo_number: str | None,
    severity: str,
    error_type: str,
    message: str,
    **extra: Any,
) -> dict[str, Any]:
    row = {
        "file": Path(file).name,
        "wo_number": wo_number,
        "severity": severity,
        "error_type": error_type,
        "message": message,
        "serial_number": extra.pop("serial_number", ""),
        "word_part": extra.pop("word_part", ""),
        "db_part": extra.pop("db_part", ""),
        "word_qty": extra.pop("word_qty", None),
        "serial_count": extra.pop("serial_count", None),
        "row_index": extra.pop("row_index", None),
    }
    row.update(extra)
    return row


def validate_word_file(file_path: Path | str) -> pd.DataFrame:
    wo_number = extract_wo_number(file_path)
    file_name = Path(file_path).name
    results = []
    items = extract_product_details(file_path)

    if not items:
        return pd.DataFrame([make_issue(
            file_name,
            wo_number,
            "WARNING",
            "NO_WO_TABLE_ROWS",
            "No readable product rows found in the first Word table.",
        )])

    seen_serials = {}
    for item in items:
        row_index = item.get("row_index")
        word_part = item["product_number"]
        word_key = normalize_part(word_part)
        serials = item["serials"]
        expected_qty_raw = item["qty"]
        expected_qty = pd.to_numeric(expected_qty_raw, errors="coerce")
        expected_qty = None if pd.isna(expected_qty) else int(expected_qty)

        if item.get("row_warning"):
            results.append(make_issue(file_name, wo_number, "WARNING", "WO_TABLE_FORMAT", item["row_warning"], word_part=word_part, row_index=row_index))

        if expected_qty is None:
            results.append(make_issue(
                file_name,
                wo_number,
                "WARNING",
                "SUSPICIOUS_QTY",
                f"Quantity is blank or non-numeric: {expected_qty_raw!r}",
                word_part=word_part,
                row_index=row_index,
            ))
        elif serials and expected_qty != len(serials):
            results.append(make_issue(
                file_name,
                wo_number,
                "CRITICAL",
                "QTY_MISMATCH",
                f"Quantity mismatch: Word {expected_qty}, serial count {len(serials)}",
                word_part=word_part,
                word_qty=expected_qty,
                serial_count=len(serials),
                row_index=row_index,
            ))

        if not serials:
            results.append(make_issue(
                file_name,
                wo_number,
                "WARNING",
                "NO_SN",
                f"No serial numbers listed; Word qty {expected_qty}",
                word_part=word_part,
                word_qty=expected_qty,
                serial_count=0,
                row_index=row_index,
            ))
            continue

        for serial in serials:
            if serial in seen_serials:
                results.append(make_issue(
                    file_name,
                    wo_number,
                    "WARNING",
                    "DUPLICATE_SN_IN_WO",
                    f"Serial number appears more than once in this WO: {serial}",
                    serial_number=serial,
                    word_part=word_part,
                    row_index=row_index,
                ))
            seen_serials[serial] = row_index

            db_part = db_part_for_serial(serial)
            if not db_part:
                results.append(make_issue(
                    file_name,
                    wo_number,
                    "WARNING",
                    "SERIAL_NOT_FOUND",
                    f"Serial number not found in receiving log: {serial}",
                    serial_number=serial,
                    word_part=word_part,
                    row_index=row_index,
                ))
            elif normalize_part(db_part) == word_key:
                results.append(make_issue(
                    file_name,
                    wo_number,
                    "INFO",
                    "MATCH",
                    "Serial part matches receiving log.",
                    serial_number=serial,
                    word_part=word_part,
                    db_part=db_part,
                    word_qty=expected_qty,
                    serial_count=len(serials),
                    row_index=row_index,
                ))
            else:
                results.append(make_issue(
                    file_name,
                    wo_number,
                    "CRITICAL",
                    "PART_MISMATCH",
                    f"Part mismatch: Word {word_part}, DB {db_part}",
                    serial_number=serial,
                    word_part=word_part,
                    db_part=db_part,
                    word_qty=expected_qty,
                    serial_count=len(serials),
                    row_index=row_index,
                ))

    return pd.DataFrame(results)


def add_cross_file_duplicate_warnings(all_results: dict[str, dict[str, Any]]) -> None:
    serial_locations = {}
    for file, result in all_results.items():
        sn_results = result.get("serial_validation")
        if sn_results is None or sn_results.empty or "serial_number" not in sn_results.columns:
            continue
        for serial in sn_results[sn_results["serial_number"].astype(str) != ""]["serial_number"].dropna().unique():
            serial_locations.setdefault(serial, set()).add(file)

    for serial, files in serial_locations.items():
        if len(files) < 2:
            continue
        file_list = ", ".join(sorted(files))
        for file in files:
            warning = make_issue(
                file,
                all_results[file].get("wo_number"),
                "WARNING",
                "DUPLICATE_SN_ACROSS_WO",
                f"Serial number appears in multiple checked WOs: {serial} ({file_list})",
                serial_number=serial,
            )
            all_results[file]["serial_validation"] = pd.concat(
                [all_results[file]["serial_validation"], pd.DataFrame([warning])],
                ignore_index=True,
            )


def report_issue_rows(all_results: dict[str, dict[str, Any]]) -> pd.DataFrame:
    rows = []
    for result in all_results.values():
        row_count = result.get("sales_order_row_count")
        if row_count and row_count["severity"] != "INFO":
            rows.append(row_count)

        sn_results = result.get("serial_validation")
        if sn_results is not None and not sn_results.empty:
            rows.extend(sn_results[sn_results["severity"].isin(["CRITICAL", "WARNING"])].to_dict("records"))
    return pd.DataFrame(rows)


def print_issue_report(
    all_results: dict[str, dict[str, Any]],
    issue_df: pd.DataFrame,
    target_date: str | datetime | None = None,
    show_warnings: bool = False,
) -> None:
    checked_count = len(all_results)
    if issue_df.empty:
        critical_df = pd.DataFrame()
        warning_df = pd.DataFrame()
    else:
        critical_df = issue_df[issue_df["severity"] == "CRITICAL"].copy()
        warning_df = issue_df[issue_df["severity"] == "WARNING"].copy()

    files_with_critical = critical_df["file"].nunique() if not critical_df.empty else 0
    clean_count = checked_count - files_with_critical

    print("TODAY'S WO CHECK")
    if target_date is not None:
        print(f"Date checked: {target_date}")
    print(f"Checked WO: {checked_count} | Critical: {len(critical_df)} | Warnings: {len(warning_df)} | WO with critical errors: {files_with_critical} | Clean WO: {clean_count}")
    print()

    print(f"{RED_ALERT} CRITICAL ISSUES {RED_ALERT}")
    if critical_df.empty:
        print("All checked WOs passed critical checks.")
    else:
        for index, file in enumerate(critical_df["file"].drop_duplicates(), start=1):
            file_errors = critical_df[critical_df["file"] == file]
            print(f"{index}. {file}")

            row_errors = file_errors[file_errors["error_type"] == "ROW_COUNT_MISMATCH"]
            for _, r in row_errors.iterrows():
                print(f"   - {RED_ALERT} Row count mismatch: Word {r.get('word_count')} | Open Sales Order {r.get('sheet_count')}")
                for mismatch in r.get("item_mismatches", []) or []:
                    if mismatch.get("mismatch_type") == "MISSING_IN_SHEET":
                        print(f"     {YELLOW_ALERT} Missing from Open Sales Order x{mismatch.get('count')}: Word {mismatch.get('word_part')}")
                    else:
                        print(f"     {YELLOW_ALERT} Extra Open Sales Order item x{mismatch.get('count')}: {mismatch.get('sheet_part')}")

            item_errors = file_errors[file_errors["error_type"] == "ITEM_NAME_MISMATCH"]
            for _, r in item_errors.iterrows():
                for mismatch in r.get("item_mismatches", []) or []:
                    if mismatch.get("mismatch_type") == "MISSING_IN_SHEET":
                        print(f"   - {YELLOW_ALERT} Missing from Open Sales Order x{mismatch.get('count')}: Word {mismatch.get('word_part')}")
                    else:
                        print(f"   - {YELLOW_ALERT} Extra Open Sales Order item x{mismatch.get('count')}: {mismatch.get('sheet_part')}")

            qty_errors = file_errors[file_errors["error_type"] == "QTY_MISMATCH"]
            for _, r in qty_errors.iterrows():
                print(f"   - {RED_ALERT} Quantity mismatch: {r.get('word_part')} | Word qty {r.get('word_qty')} | SN count {r.get('serial_count')}")

            if {"word_part", "db_part", "serial_number"}.issubset(file_errors.columns):
                part_errors = file_errors[file_errors["error_type"] == "PART_MISMATCH"]
                for (word_part, db_part), group in part_errors.groupby(["word_part", "db_part"], dropna=False):
                    serials = [str(sn) for sn in group["serial_number"].dropna().tolist() if str(sn)]
                    shown = ", ".join(serials[:6])
                    more = f" (+{len(serials) - 6} more)" if len(serials) > 6 else ""
                    print(f"   - {RED_ALERT} Part mismatch x{len(group)}: {shown}{more} | Word {word_part} | DB {db_part}")
        print()

    print("WARNINGS")
    if warning_df.empty:
        print("No warnings.")
        return

    warning_summary = warning_df.groupby("error_type").size().sort_values(ascending=False)
    summary_text = "; ".join(f"{name}: {count}" for name, count in warning_summary.items())
    print(f"{len(warning_df)} warnings hidden. {summary_text}")

    if show_warnings:
        print()
        for index, (_, r) in enumerate(warning_df.iterrows(), start=1):
            serial = f" | SN {r.get('serial_number')}" if r.get("serial_number") else ""
            part = f" | Part {r.get('word_part')}" if r.get("word_part") else ""
            print(f"{index}. {r.get('file')}{serial}{part} | {r.get('message')}")
    else:
        print("Use --show-warnings to print warning details.")


def validate_wo_folder(
    folder: Path | str = WO_FOLDER,
    sales_order: pd.DataFrame | None = None,
    days: int = 0,
    show_warnings: bool = False,
) -> dict[str, Any]:
    target_date = datetime.today().date() - timedelta(days)
    all_results: dict[str, dict[str, Any]] = {}

    for root, _, files in os.walk(folder):
        for file in files:
            if not file.lower().endswith(".docx"):
                continue

            file_path = os.path.join(root, file)
            if not is_real_wo_file(file_path):
                continue

            creation_time = datetime.fromtimestamp(os.path.getctime(file_path)).date()
            modified_time = datetime.fromtimestamp(os.path.getmtime(file_path)).date()
            if creation_time != target_date and modified_time != target_date:
                continue

            sn_results = validate_word_file(file_path)
            row_count_result = compare_word_file_to_sales_order_result(file_path, sales_order=sales_order)
            all_results[file] = {
                "wo_number": extract_wo_number(file_path),
                "file_path": file_path,
                "created_date": creation_time,
                "modified_date": modified_time,
                "serial_validation": sn_results,
                "sales_order_row_count": row_count_result,
                "sales_order_count_check": compare_word_file_to_sales_order(file_path, sales_order=sales_order),
            }

            if creation_time != modified_time and modified_time == target_date:
                warning = make_issue(
                    file,
                    extract_wo_number(file_path),
                    "WARNING",
                    "OLD_FILE_MODIFIED_TODAY",
                    f"File was created {creation_time} and modified {modified_time}.",
                )
                all_results[file]["serial_validation"] = pd.concat([sn_results, pd.DataFrame([warning])], ignore_index=True)

    add_cross_file_duplicate_warnings(all_results)
    issue_df = report_issue_rows(all_results)
    critical_df = issue_df[issue_df["severity"] == "CRITICAL"].copy() if not issue_df.empty else pd.DataFrame()
    warning_df = issue_df[issue_df["severity"] == "WARNING"].copy() if not issue_df.empty else pd.DataFrame()

    report_results: dict[str, Any] = dict(all_results)
    report_results["_report"] = {
        "issues": issue_df,
        "critical": critical_df,
        "warnings": warning_df,
    }
    print_issue_report(all_results, issue_df, target_date=target_date, show_warnings=show_warnings)
    return report_results


def validate_single_file(
    file_path: Path | str,
    sales_order: pd.DataFrame | None = None,
    show_warnings: bool = False,
) -> dict[str, Any]:
    file_path = Path(file_path)
    sn_results = validate_word_file(file_path)
    row_count_result = compare_word_file_to_sales_order_result(file_path, sales_order=sales_order)
    all_results = {
        file_path.name: {
            "wo_number": extract_wo_number(file_path),
            "file_path": str(file_path),
            "serial_validation": sn_results,
            "sales_order_row_count": row_count_result,
            "sales_order_count_check": compare_word_file_to_sales_order(file_path, sales_order=sales_order),
        }
    }
    issue_df = report_issue_rows(all_results)
    critical_df = issue_df[issue_df["severity"] == "CRITICAL"].copy() if not issue_df.empty else pd.DataFrame()
    warning_df = issue_df[issue_df["severity"] == "WARNING"].copy() if not issue_df.empty else pd.DataFrame()
    all_results["_report"] = {"issues": issue_df, "critical": critical_df, "warnings": warning_df}
    print_issue_report({file_path.name: all_results[file_path.name]}, issue_df, show_warnings=show_warnings)
    return all_results


def export_report(results: dict[str, Any], output_dir: Path | str) -> dict[str, Path]:
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    report = results.get("_report", {})
    exported = {}
    for name in ["issues", "critical", "warnings"]:
        df = report.get(name)
        if isinstance(df, pd.DataFrame):
            path = output_dir / f"wo_{name}.csv"
            df.to_csv(path, index=False)
            exported[name] = path
    return exported
